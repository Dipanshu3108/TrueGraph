"""DOCX parser."""

import re
from typing import Any, Dict, List, Optional

from ..cleanup.text_cleanup import cleanup_page
from ..exceptions import DependencyError, ParseError
from ..logger import get_logger

logger = get_logger()

try:
    import docx
except ImportError as exc:  # pragma: no cover
    docx = None  # type: ignore


def _has_page_break(paragraph: Any) -> bool:
    """Detect manual page breaks in a paragraph."""
    for run in paragraph.runs:
        if run._element.xpath(".//w:br[@w:type='page']"):
            return True
    return False


class DOCXParser:
    """Parse DOCX files including paragraphs, tables, headers, footers, and links."""

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.config = config or {}

    def parse(self, path: str) -> List[Dict[str, Any]]:
        if docx is None:
            raise DependencyError(
                "python-docx is required. Install it with: pip install python-docx"
            )

        try:
            document = docx.Document(path)
        except Exception as exc:
            raise ParseError(f"Failed to open DOCX {path}: {exc}") from exc

        chunks: List[str] = []

        # Header(s)
        for section in document.sections:
            for header in section.header, section.first_page_header, section.even_page_header:
                if header is not None:
                    for para in header.paragraphs:
                        text = para.text.strip()
                        if text:
                            chunks.append(text)

        # Body paragraphs
        for para in document.paragraphs:
            if _has_page_break(para):
                chunks.append("\f")
            text = para.text.strip()
            if text:
                chunks.append(text)

        # Tables
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    chunks.append(" | ".join(cells))

        # Footer(s)
        for section in document.sections:
            for footer in section.footer, section.first_page_footer, section.even_page_footer:
                if footer is not None:
                    for para in footer.paragraphs:
                        text = para.text.strip()
                        if text:
                            chunks.append(text)

        return self._split_into_pages(chunks)

    def _split_into_pages(self, chunks: List[str]) -> List[Dict[str, Any]]:
        """Split collected text chunks into pages using form-feed markers."""
        pages: List[Dict[str, Any]] = []
        current: List[str] = []
        page_no = 1

        for chunk in chunks:
            if chunk == "\f":
                content = cleanup_page("\n".join(current))
                if content:
                    pages.append({"page_no": page_no, "content": content})
                    page_no += 1
                current = []
            else:
                current.append(chunk)

        if current:
            content = cleanup_page("\n".join(current))
            if content:
                pages.append({"page_no": page_no, "content": content})

        # If nothing produced pages, return a single page with all content.
        if not pages:
            all_text = cleanup_page("\n".join(chunks))
            pages.append({"page_no": 1, "content": all_text})

        return pages
