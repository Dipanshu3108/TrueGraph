"""Tests for DOCXParser."""

import os
import tempfile
import unittest

from libs.UDP.universal_document_parser.parsers.docx_parser import DOCXParser


try:
    import docx

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


@unittest.skipUnless(HAS_DOCX, "python-docx not installed")
class TestDOCXParser(unittest.TestCase):
    def test_parse_docx(self):
        from docx import Document

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()
            doc.add_paragraph("First paragraph.")
            doc.add_paragraph("Second paragraph.")
            doc.save(tmp.name)
            tmp_path = tmp.name

        try:
            parser = DOCXParser()
            result = parser.parse(tmp_path)
            self.assertTrue(len(result) >= 1)
            content = result[0]["content"]
            self.assertIn("First paragraph.", content)
            self.assertIn("Second paragraph.", content)
        finally:
            os.unlink(tmp_path)


if __name__ == "__main__":
    unittest.main()
